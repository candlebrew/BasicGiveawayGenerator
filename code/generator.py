import random
import csv
import getpass
import datetime
import docx

totalEntries = 0
timestamp = ""
winnerName = "" # our winner's name
winnerID = "" # our winner's unique ID number
winnerEmail = "" # our winner's contact email
filename = "entries.csv"

currentUser = getpass.getuser()

def enterFile():
    global filename
    try:
        with open(filename, newline="") as openfile:
            pass
        return filename
    except FileNotFoundError:
        while True:
            print("What is the name of the CSV file containing giveaway entrants?")
            filename = input("")
            if ".csv" not in filename:
                filename += ".csv"
            try:
                with open(filename, newline="") as openfile:
                    pass
                return filename
            except FileNotFoundError:
                print(f"I don't recognize the filename '{filename}' in the local directory. Please try again.")

def generate():
    global filename, totalEntries, timestamp, winnerName, winnerID, winnerEmail
    filename = enterFile()

    now = datetime.datetime.now()
    timestamp = now.strftime("%B %d, %Y %I:%M:%S %p")

    with open(filename, newline="") as entriesCSV:
        entriesDict = csv.reader(entriesCSV,dialect="excel")

        totalEntries = len(list(entriesDict)) - 1 # ignore our header row

    winningNumber = random.randint(1,totalEntries)

    with open(filename, newline="") as entriesCSV:
        entriesDict = csv.DictReader(entriesCSV,dialect="excel")

        for row in entriesDict:
            if int(row["#"]) == winningNumber:
                winnerName = row["Name"]
                winnerID = row["ID"]
                winnerEmail = row["Email"]
                print(f"The winner is {winnerName}, ID {winnerID}, email {winnerEmail}")

def createAuditSheet():
    doc = docx.Document()
    doc.add_paragraph("Giveaway: __________________________________________________________________________________________")
    doc.add_paragraph("______________________________________________________________________________________________________")
    doc.add_paragraph("Prize: _______________________________________________________________________________________________")
    doc.add_paragraph(f"The winner is {winnerName}, ID {winnerID}, email {winnerEmail}.")
    doc.add_paragraph(f"Drawn {timestamp} by user {currentUser} from {totalEntries} total entries found in file {filename}")
    for i in range(5):
        doc.add_paragraph("")
    doc.add_paragraph("Signature _________________________________________________________________ Date ___________________")
    doc.add_paragraph("Pick Up Date ___________________")
    doc.save("GiveawayDrawingResults.docx")

generate()
createAuditSheet()
print("Task completed.")
